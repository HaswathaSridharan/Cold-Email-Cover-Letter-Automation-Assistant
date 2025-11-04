# ================== main.py (Full Agentic Version) ==================

import os
import streamlit as st
from dotenv import load_dotenv
from chains import Chain, IndividualChain
from portfolio import Portfolio
from utils import clean_text
from langchain_community.document_loaders import WebBaseLoader
import pdfplumber
from docx import Document
from docx import Document as DocxDocument
from io import BytesIO
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import base64
import pickle
import datetime

load_dotenv()

# ----------------- GMAIL AUTHENTICATION -----------------
def authenticate_gmail():
    SCOPES = ['https://mail.google.com/']
    creds = None

    current_dir = os.path.dirname(os.path.abspath(__file__))
    credentials_path = os.path.join(current_dir, 'credentials.json')
    token_path = os.path.join(current_dir, "token.pickle")

    if os.path.exists(token_path):
        with open(token_path, "rb") as token:
            creds = pickle.load(token)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
            creds = flow.run_local_server(port=8080)
        with open(token_path, "wb") as token:
            pickle.dump(creds, token)

    service = build('gmail', 'v1', credentials=creds)
    profile = service.users().getProfile(userId='me').execute()
    from_email = profile['emailAddress']
    return service, from_email

# ----------------- SEND EMAIL WITH ATTACHMENTS -----------------
def send_email_with_attachments(service, from_email, to_email, subject, body_text, resume_file, cover_letter_text):
    message = MIMEMultipart()
    message['to'] = to_email
    message['from'] = from_email
    message['subject'] = subject

    # Attach body
    message.attach(MIMEText(body_text, 'plain'))

    # Attach resume
    if resume_file:
        resume_part = MIMEBase('application', 'octet-stream')
        resume_part.set_payload(resume_file.read())
        encoders.encode_base64(resume_part)
        resume_part.add_header('Content-Disposition', f'attachment; filename="{resume_file.name}"')
        message.attach(resume_part)

    # Attach generated cover letter as a .docx file
    cover_letter_doc = DocxDocument()
    for line in cover_letter_text.split('\n'):
        cover_letter_doc.add_paragraph(line)
    cover_letter_io = BytesIO()
    cover_letter_doc.save(cover_letter_io)
    cover_letter_io.seek(0)

    cover_letter_part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
    cover_letter_part.set_payload(cover_letter_io.read())
    encoders.encode_base64(cover_letter_part)
    cover_letter_part.add_header('Content-Disposition', 'attachment', filename="Cover_Letter.docx")
    message.attach(cover_letter_part)

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    body = {'raw': raw_message}

    sent_message = service.users().messages().send(userId="me", body=body).execute()
    return sent_message

# ----------------- PAGE NAVIGATION -----------------
def go_to_individual():
    st.session_state.page = 'individual'

def go_to_organization():
    st.session_state.page = 'organization'

def go_home():
    st.session_state.page = 'landing'

# ----------------- MAIN APP -----------------
if 'page' not in st.session_state:
    st.session_state.page = 'landing'

if st.session_state.page == 'landing':
    st.set_page_config(page_title="Cold Email Assistant", page_icon="📧")
    st.title("📧 Welcome to the Cold Email Assistant!")

    col1, col2 = st.columns(2)
    with col1:
        st.button("🙋 Individual", on_click=go_to_individual)
    with col2:
        st.button("🏢 Organization", on_click=go_to_organization)

# ----------------- INDIVIDUAL PAGE -----------------
elif st.session_state.page == 'individual':
    st.title("🚀 Cold Email Generator for Individuals")

    chain = IndividualChain()
    action_log = []

    uploaded_resume = st.file_uploader("Upload your Resume (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
    job_url = st.text_input("Paste a Job URL here")

    if st.button("Generate Cold Email", disabled=not (uploaded_resume and job_url)):
        try:
            action_log.append("Loading and cleaning job page.")
            loader = WebBaseLoader([job_url])
            page_data = loader.load().pop().page_content
            cleaned_data = clean_text(page_data)

            action_log.append("Extracting job information using LLM.")
            job = chain.extract_job_info(cleaned_data)

            resume_text = ""
            if uploaded_resume.type == "application/pdf":
                with pdfplumber.open(uploaded_resume) as pdf:
                    for page in pdf.pages:
                        resume_text += page.extract_text()
            elif uploaded_resume.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                doc = Document(uploaded_resume)
                resume_text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_resume.type == "text/plain":
                resume_text = uploaded_resume.read().decode('utf-8')

            action_log.append("Extracting applicant information from resume.")
            applicant_info = chain.extract_contact_info(resume_text)

            recruiter_email = job.get("recruiter_email")
            if not recruiter_email:
                action_log.append("Primary recruiter email not found, attempting fallback lookup.")
                recruiter_email = chain.lookup_recruiter_email(job.get("recruiter_name", ""), job.get("company_name", ""))
            st.session_state.hunter_recruiter_email = recruiter_email or "hr@company.com"  # fallback if lookup fails

            action_log.append("Generating cover letter and cold email.")
            cover_letter = chain.generate_cover_letter(job, resume_text)
            cold_email = chain.generate_cold_email(job, applicant_info, cover_letter)

            st.session_state.generated_email = cold_email
            st.session_state.generated_cover_letter = cover_letter
            st.session_state.email_ready = True

            st.subheader("📄 Generated Cover Letter")
            st.code(cover_letter, language='markdown')

            st.subheader("📧 Generated Cold Email")
            editable_email = st.text_area("📝 Edit Cold Email Before Sending:", value=cold_email, height=300)
            st.session_state.editable_email = editable_email

            st.success("✅ Cold Email Generated Successfully!")
            st.session_state.action_log = action_log

        except Exception as e:
            st.error(f"❌ Error occurred during generation: {e}")
            action_log.append(f"Error: {e}")

    if st.session_state.get("email_ready", False):
        st.divider()
        st.info(f"*Recruiter's Email (via Hunter.io):* {st.session_state.hunter_recruiter_email or 'Not Found'}")

        final_receiver_email = st.text_input("📨 Enter the Email Address to whom the mail should be sent:")
        your_email = st.text_input("✉ Enter Your Own Email Address (From Address):")

        if st.button("📤 Confirm and Send Email Now"):
            try:
                if not final_receiver_email:
                    st.error("❗ Please enter recipient's Email Address.")
                elif not your_email:
                    st.error("❗ Please enter your Email Address.")
                else:
                    service, from_email = authenticate_gmail()

                    send_email_with_attachments(
                        service=service,
                        from_email=your_email,
                        to_email=final_receiver_email,
                        subject="Exciting Career Opportunity Inquiry",
                        body_text=st.session_state.editable_email,
                        resume_file=uploaded_resume,
                        cover_letter_text=st.session_state.generated_cover_letter
                    )

                    st.success(f"✅ Email with attachments sent to {final_receiver_email}!")
                    st.balloons()

                    

            except Exception as e:
                st.error(f"❌ Error sending email: {e}")

    st.button("⬅ Back", on_click=go_home)

# (Organization page remains same, can be similarly upgraded)
# ----------------- ORGANIZATION PAGE -----------------
elif st.session_state.page == 'organization':
    st.title("🏢 Cold Email Generator for Organizations")

    chain = Chain()
    portfolio = Portfolio(file_path=r"/Users/juhianand/Documents/UIC/Spring/DeepLearning/Cold_Email_Project/app/resource/my_portfolio.csv")
    action_log = []

    url_input = st.text_input("🌐 Enter Organization Careers Page URL:")

    if st.button("🔍 Generate Organization Cold Email", disabled=not url_input):
        try:
            action_log.append("Loading and cleaning organization's careers page.")
            loader = WebBaseLoader([url_input])
            page_data = loader.load().pop().page_content
            cleaned_data = clean_text(page_data)

            action_log.append("Extracting job descriptions using LLM.")
            jobs = chain.extract_jobs(cleaned_data)

            if jobs:
                job = jobs[0]  # For now, take the first job found
                skills = job.get("skills", [])
                links = portfolio.query_links(skills)

                action_log.append("Generating cold email for organization.")
                cold_email = chain.write_mail(job, links)

                # Recruiter Email Handling
                recruiter_name = job.get("recruiter_name", "")
                company_name = job.get("company_name", "")

                hunter_chain = IndividualChain()
                hunter_email = ""
                if recruiter_name and company_name:
                    hunter_email = hunter_chain.lookup_recruiter_email(recruiter_name, company_name)
                st.session_state.hunter_recruiter_email_org = hunter_email or "hr@company.com"

                st.session_state.generated_org_email = cold_email
                st.session_state.org_email_ready = True

                st.subheader("📧 Generated Organization Cold Email")
                editable_org_email = st.text_area("📝 Edit Organization Cold Email Before Sending:", value=cold_email, height=300)
                st.session_state.editable_org_email = editable_org_email

                st.success("✅ Organization Cold Email Generated!")
                st.session_state.org_action_log = action_log

            else:
                st.warning("⚠ No jobs found from the careers page.")
                action_log.append("No jobs extracted.")

        except Exception as e:
            st.error(f"❌ Error occurred during organization cold email generation: {e}")
            action_log.append(f"Error: {e}")

    if st.session_state.get("org_email_ready", False):
        st.divider()
        st.info(f"*Recruiter's Email (via Hunter.io):* {st.session_state.hunter_recruiter_email_org or 'Not Found'}")

        final_receiver_email = st.text_input("📨 Enter the Organization Email Address to send to:")
        your_email = st.text_input("✉ Enter Your Own Email Address (From Address):")

        if st.button("📤 Confirm and Send Organization Email"):
            try:
                if not final_receiver_email:
                    st.error("❗ Please enter recipient's Email Address.")
                elif not your_email:
                    st.error("❗ Please enter your Email Address.")
                else:
                    service, from_email = authenticate_gmail()

                    # Now sending simple text email without attachments
                    message = MIMEText(st.session_state.editable_org_email)
                    message['to'] = final_receiver_email
                    message['from'] = your_email
                    message['subject'] = "Exciting Collaboration Opportunity with InnovaEdge Technologies"

                    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()
                    body = {'raw': raw}

                    sent_message = service.users().messages().send(userId="me", body=body).execute()

                    if sent_message and sent_message.get("id"):
                        st.success(f"✅ Organization Email sent to {final_receiver_email}!")
                        st.balloons()
                        
                    else:
                        st.error("❌ Failed to send Organization Email.")

            except Exception as e:
                st.error(f"❌ Error sending Organization email: {e}")

    st.button("⬅ Back", on_click=go_home)
